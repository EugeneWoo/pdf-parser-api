import os
import base64
import io
import re
from datetime import datetime
from docx import Document
from flask import Blueprint, request, jsonify

fill_retainer_bp = Blueprint("fill_retainer", __name__)

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "retainer_template.docx")


def normalize_date(raw: str) -> str:
    """Coerce any date-like string to YYYY-MM-DD."""
    if not raw:
        return raw
    # Already YYYY-MM-DD
    if re.match(r"^\d{4}-\d{2}-\d{2}$", raw):
        return raw
    # ISO 8601 timestamp e.g. 2018-12-06T05:00:00.000Z
    try:
        return datetime.fromisoformat(raw.replace("Z", "+00:00")).strftime("%Y-%m-%d")
    except (ValueError, AttributeError):
        return raw


def build_injury_paragraph(data, num_injured):
    """Build the injury paragraph based on whether there were injuries."""
    if num_injured > 0:
        return ("Additionally, since the motor vehicle accident involved an injured person, Attorney will "
                "also investigate potential bodily injury claims and review relevant medical records "
                "to substantiate non-economic damages.")
    else:
        return ("However, since the motor vehicle accident involved no reported injured people, the scope "
                "of this engagement is strictly limited to the recovery of property damage and loss of use.")


@fill_retainer_bp.route("/fill-retainer", methods=["POST"])
def fill_retainer():
    data = request.get_json(silent=True)
    if not data:
        return jsonify({"error": "JSON body required"}), 400

    # Accept either direct data or data wrapped in "data" key from /parse endpoint
    parse_data = data.get("data", data)

    required = ["client_first_name", "client_last_name", "Accident Date", "Defendant Name",
                "Client Gender", "Accident Location", "Client Plate Number",
                "Accident Description", "Num Injured", "Statute of Limitations Date"]
    missing = [k for k in required if k not in parse_data]
    if missing:
        return jsonify({"error": f"Missing fields: {missing}"}), 400

    if not os.path.exists(TEMPLATE_PATH):
        return jsonify({"error": "Template file not found on server"}), 500

    doc = Document(TEMPLATE_PATH)

    def replace_in_paragraph(para, replacements):
        """Replace placeholders that may span multiple runs."""
        for key, val in replacements.items():
            placeholder = "{{" + key + "}}"
            if placeholder not in para.text:
                continue

            # Replace in the full paragraph text
            para.text = para.text.replace(placeholder, val)

    # Transform data from /parse format to template format
    gender = parse_data.get("Client Gender", "M").upper()
    num_injured = parse_data.get("Num Injured", 0)

    replacements = {
        "client_name": f"{parse_data['client_first_name']} {parse_data['client_last_name']}",
        "accident_date": normalize_date(parse_data["Accident Date"]),
        "defendant_name": parse_data["Defendant Name"],
        "pronoun_his_her": "his" if gender == "M" else "her",
        "pronoun_he_she": "he" if gender == "M" else "she",
        "pronoun_him_her": "him" if gender == "M" else "her",
        "accident_location": parse_data["Accident Location"],
        "client_plate": parse_data["Client Plate Number"],
        "accident_description": parse_data["Accident Description"],
        "injury_paragraph": build_injury_paragraph(parse_data, num_injured),
        "sol_date": normalize_date(parse_data["Statute of Limitations Date"]),
    }

    for para in doc.paragraphs:
        replace_in_paragraph(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, replacements)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    docx_base64 = base64.standard_b64encode(buf.read()).decode("utf-8")
    return jsonify({"docx_base64": docx_base64})
