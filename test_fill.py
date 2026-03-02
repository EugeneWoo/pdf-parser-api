import os
import sys
import base64
import io

# Add the current directory to path so we can import the module
sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
from fill_retainer import build_injury_paragraph

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "retainer_template.docx")

# Test data matching /parse endpoint output
test_data = {
    "client_first_name": "Guillermo",
    "client_last_name": "REYES",
    "Accident Date": "2018-12-06",
    "Defendant Name": "FRANCOIS, LIONEL",
    "Client Gender": "M",
    "Accident Location": "Flatbush Avenue and Plaza Street East, Kings County, NY",
    "Client Plate Number": "XCGY85",
    "Accident Description": "V1 was driving northbound on Flatbush Avenue in the right lane when V2 came over from the bus lane into the middle lane and struck V1. V2 then came over and struck V1 again with no injuries; officer did not witness the accident.",
    "Num Injured": 0,
    "Statute of Limitations Date": "2026-12-06"
}

print("=== Test Data ===")
print(f"Client: {test_data['client_first_name']} {test_data['client_last_name']}")
print(f"Gender: {test_data['Client Gender']}")
print(f"Num Injured: {test_data['Num Injured']}")
print()

# Build the injury paragraph
injury_paragraph = build_injury_paragraph(test_data, test_data['Num Injured'])
print("=== Injury Paragraph ===")
print(injury_paragraph)
print()

# Test the template fill
if not os.path.exists(TEMPLATE_PATH):
    print(f"ERROR: Template not found at {TEMPLATE_PATH}")
    sys.exit(1)

doc = Document(TEMPLATE_PATH)

# Build replacements
gender = test_data.get("Client Gender", "M").upper()
num_injured = test_data.get("Num Injured", 0)

replacements = {
    "client_name": f"{test_data['client_first_name']} {test_data['client_last_name']}",
    "accident_date": test_data["Accident Date"],
    "defendant_name": test_data["Defendant Name"],
    "pronoun_his_her": "his" if gender == "M" else "her",
    "pronoun_he_she": "he" if gender == "M" else "she",
    "pronoun_him_her": "him" if gender == "M" else "her",
    "accident_location": test_data["Accident Location"],
    "client_plate": test_data["Client Plate Number"],
    "accident_description": test_data["Accident Description"],
    "injury_paragraph": injury_paragraph,
    "sol_date": test_data["Statute of Limitations Date"],
}

print("=== Replacements ===")
for key, value in replacements.items():
    print(f"{{{{{key}}}}} → {value}")
print()

# Replace in document
def replace_in_paragraph(para, replacements):
    """Replace placeholders that may span multiple runs."""
    replaced_count = 0
    for key, val in replacements.items():
        placeholder = "{{" + key + "}}"
        if placeholder in para.text:
            para.text = para.text.replace(placeholder, val)
            replaced_count += 1
    return replaced_count > 0

found_placeholders = []
for para in doc.paragraphs:
    for key in replacements.keys():
        placeholder = "{{" + key + "}}"
        if placeholder in para.text:
            if key not in found_placeholders:
                found_placeholders.append(key)
            replace_in_paragraph(para, replacements)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for key in replacements.keys():
                    placeholder = "{{" + key + "}}"
                    if placeholder in para.text:
                        if key not in found_placeholders:
                            found_placeholders.append(key)
                        replace_in_paragraph(para, replacements)

# Check for missing placeholders
missing_placeholders = set(replacements.keys()) - set(found_placeholders)
if missing_placeholders:
    print("=== WARNING: Placeholders NOT found in template ===")
    for key in missing_placeholders:
        print(f"  {{{{{key}}}}}")
else:
    print("=== SUCCESS: All placeholders found and replaced ===")

print()
print("=== Placeholders Found ===")
for key in found_placeholders:
    print(f"  {{{{{key}}}}}")

# Save the test document
output_path = os.path.join(os.path.dirname(__file__), "test_output.docx")
doc.save(output_path)
print()
print(f"Test document saved to: {output_path}")
